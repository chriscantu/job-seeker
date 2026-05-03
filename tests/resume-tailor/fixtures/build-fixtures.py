#!/usr/bin/env python3
"""
Generate minimal .docx test fixtures for page-count integration tests.
Idempotent — safe to re-run; overwrites existing files.

Outputs:
  one-page.docx   — single paragraph; soffice reports 1 page
  three-page.docx — three paragraphs separated by page breaks; soffice reports 3 pages

Usage:
  python3 tests/resume-tailor/fixtures/build-fixtures.py
"""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent

# ── one-page.docx ──────────────────────────────────────────────────────────────
doc = Document()
doc.add_paragraph("one page fixture")
doc.save(HERE / "one-page.docx")
print("wrote one-page.docx")

# ── three-page.docx ────────────────────────────────────────────────────────────
doc = Document()
doc.add_paragraph("page one")

# explicit page break via run
p1 = doc.add_paragraph()
run = p1.add_run()
br = OxmlElement("w:br")
br.set(qn("w:type"), "page")
run._r.append(br)

doc.add_paragraph("page two")

p2 = doc.add_paragraph()
run2 = p2.add_run()
br2 = OxmlElement("w:br")
br2.set(qn("w:type"), "page")
run2._r.append(br2)

doc.add_paragraph("page three")

doc.save(HERE / "three-page.docx")
print("wrote three-page.docx")
