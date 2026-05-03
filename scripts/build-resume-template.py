#!/usr/bin/env python3
"""
build-resume-template.py
========================
Generates references/resume-template.docx — a Word document containing ONLY
named paragraph styles (no body content). The template is consumed by the
anthropic-skills:docx renderer at render time (Task 5.1). Style names in this
file MUST exactly match the names the renderer emits per the render contract
at docs/superpowers/specs/2026-05-01-ats-resume-template-design.md lines 257-269.

Verified working with python-docx 1.2.0.

Style inventory
---------------
Heading 1        — Candidate name (built-in override): 22pt Calibri bold navy, centered
Tagline          — Subtitle / positioning line: 12pt Calibri italic navy, centered
Contact          — Contact info line: 10pt Calibri dark-gray, centered
Heading 2        — Section headers (built-in override): 14pt Calibri bold navy all-caps,
                   0.5pt navy bottom border
Heading 3        — Role title (built-in override): 12pt Calibri bold navy
Role Meta        — Company / date / location metadata: 10pt Calibri italic gray
List Bullet      — Achievement bullets (built-in override): 11pt Calibri, bullet glyph •
Skills Line      — Pipe-separated skills row: 11pt Calibri
Accomplishment   — Impact bullets (same glyph as List Bullet): 11pt Calibri, bullet glyph •
Body Text (Summary) — Summary paragraph: 11pt Calibri justified

Note on "Body Text (Summary)": python-docx accepts parentheses in the display
name; the style ID is normalised to "BodyTextSummary" internally by Word but
the display name is preserved as authored. If a future renderer version cannot
address the style by its display name, fall back to style ID lookup.

Idempotent: running this script multiple times produces identical output because
styles are overwritten deterministically from the same constants.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants  # noqa — version guard

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "references",
    "resume-template.docx",
)

# ── Palette ────────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
DARK_GRAY = RGBColor(0x3A, 0x3A, 0x3A)
MID_GRAY = RGBColor(0x5A, 0x5A, 0x5A)
BLACK = RGBColor(0x00, 0x00, 0x00)
BULLET = "•"  # U+2022 BULLET


# ── Helpers ─────────────────────────────────────────────────────────────────

def _set_font(run_fmt, name, size_pt, bold=False, italic=False, color=None):
    """Apply font attributes to a _CharacterFormat (run or paragraph font)."""
    run_fmt.name = name
    run_fmt.size = Pt(size_pt)
    run_fmt.bold = bold
    run_fmt.italic = italic
    if color is not None:
        run_fmt.color.rgb = color


def _set_para_spacing(para_fmt, before_pt=0, after_pt=0, line_rule=None, line_val=None):
    """Set paragraph spacing. line_rule / line_val use docx WD_LINE_SPACING constants."""
    para_fmt.space_before = Pt(before_pt)
    para_fmt.space_after = Pt(after_pt)
    if line_rule is not None:
        para_fmt.line_spacing_rule = line_rule
        para_fmt.line_spacing = line_val


def _add_bottom_border(style, color_hex="1F3A5F", sz_eighths=4):
    """
    Attach a bottom paragraph border via raw OOXML.

    Takes the Style object (not paragraph_format) because for styles,
    paragraph_format._element returns the CT_Style element, not pPr.
    We access pPr directly via style.element.get_or_add_pPr().

    sz_eighths: border width in 1/8-pt units. 0.5pt = 4 eighths.
    color_hex: 6-char hex without '#'.
    """
    pPr = style.element.get_or_add_pPr()
    # pPr may already have a pBdr child; reuse or create
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz_eighths))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)


def _add_custom_style(doc, name, base_style_name=None):
    """
    Add a new custom paragraph style. Returns the Style object.
    Raises if the name already exists (caller should use doc.styles[name] instead).
    """
    from docx.enum.style import WD_STYLE_TYPE
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base_style_name:
        style.base_style = doc.styles[base_style_name]
    return style


def _get_or_add_style(doc, name, base_style_name=None):
    """Return existing style by name, or create it."""
    try:
        return doc.styles[name]
    except KeyError:
        return _add_custom_style(doc, name, base_style_name)


# ── Page setup ──────────────────────────────────────────────────────────────

def configure_page(doc):
    """Set 0.75-inch margins on all sides."""
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


# ── Style definitions ────────────────────────────────────────────────────────

def style_heading1(doc):
    """
    Heading 1 — Candidate name.
    Built-in override: 22pt Calibri bold, navy, centered, 6pt after.
    """
    s = doc.styles["Heading 1"]
    _set_font(s.font, "Calibri", 22, bold=True, color=NAVY)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(s.paragraph_format, before_pt=0, after_pt=6)
    return s


def style_heading2(doc):
    """
    Heading 2 — Section headers.
    Built-in override: 14pt Calibri bold, navy, ALL CAPS, 8pt before, 4pt after,
    0.5pt navy bottom border.
    """
    s = doc.styles["Heading 2"]
    _set_font(s.font, "Calibri", 14, bold=True, color=NAVY)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(s.paragraph_format, before_pt=8, after_pt=4)
    # ALL CAPS via font.all_caps
    s.font.all_caps = True
    _add_bottom_border(s, color_hex="1F3A5F", sz_eighths=4)
    return s


def style_heading3(doc):
    """
    Heading 3 — Role title within an experience block.
    Built-in override: 12pt Calibri bold, navy, 6pt before, 2pt after.
    """
    s = doc.styles["Heading 3"]
    _set_font(s.font, "Calibri", 12, bold=True, color=NAVY)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(s.paragraph_format, before_pt=6, after_pt=2)
    return s


def style_tagline(doc):
    """
    Tagline — Subtitle / positioning statement below candidate name.
    12pt Calibri italic, navy, centered, 0pt before, 4pt after.
    """
    s = _get_or_add_style(doc, "Tagline", base_style_name="Normal")
    _set_font(s.font, "Calibri", 12, italic=True, color=NAVY)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(s.paragraph_format, before_pt=0, after_pt=4)
    return s


def style_contact(doc):
    """
    Contact — Contact info line(s).
    10pt Calibri, dark gray, centered, 0pt before, 12pt after.
    """
    s = _get_or_add_style(doc, "Contact", base_style_name="Normal")
    _set_font(s.font, "Calibri", 10, color=DARK_GRAY)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(s.paragraph_format, before_pt=0, after_pt=12)
    return s


def style_role_meta(doc):
    """
    Role Meta — Company / dates / location beneath Heading 3.
    10pt Calibri italic, mid gray, 0pt before, 4pt after.
    """
    s = _get_or_add_style(doc, "Role Meta", base_style_name="Normal")
    _set_font(s.font, "Calibri", 10, italic=True, color=MID_GRAY)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(s.paragraph_format, before_pt=0, after_pt=4)
    return s


def style_list_bullet(doc):
    """
    List Bullet — Achievement bullets (built-in override).
    11pt Calibri, black, bullet glyph •, 0.15in hanging indent,
    2pt before, 2pt after, 1.0 line height (single).
    """
    from docx.enum.text import WD_LINE_SPACING

    s = doc.styles["List Bullet"]
    _set_font(s.font, "Calibri", 11, color=BLACK)
    pf = s.paragraph_format
    _set_para_spacing(pf, before_pt=2, after_pt=2,
                      line_rule=WD_LINE_SPACING.SINGLE, line_val=Pt(12))
    pf.left_indent = Inches(0.15)
    pf.first_line_indent = Inches(-0.15)  # hanging

    # Ensure bullet glyph is • via numPr or direct tab/bullet approach.
    # python-docx List Bullet style inherits Word's built-in numbering list.
    # We override the numFmt / numId to a simple bullet at the OOXML level.
    # Simplest reliable path: set numId=0 (disable list numbering) and rely
    # on the renderer to prepend the • glyph, OR keep the built-in list
    # numbering and ensure lvlText is set to •.
    # Strategy: override the numFmt on the abstract num definition is fragile
    # without the full numbering XML. Instead, we set the paragraph numPr
    # to reference an abstract num with bullet format that we create.
    _set_list_bullet_glyph(doc, s)
    return s


def _set_list_bullet_glyph(doc, style):
    """
    Ensure List Bullet style uses • as its bullet glyph.

    We create a minimal abstract numbering definition with a single level
    that uses the • glyph and attach it to the style. If a numbering part
    already exists we append to it; otherwise we create one.
    """
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    from lxml import etree

    # Access or create the numbering part
    try:
        num_part = doc.part.numbering_part
    except AttributeError:
        # No numbering part yet — create one
        num_xml = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        num_part = Part(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            etree.fromstring(num_xml),
            doc.part.package,
        )
        doc.part.relate_to(
            num_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
        )

    # Assign a deterministic abstract num id and num id to avoid duplicates
    ABSTRACT_ID = 100
    NUM_ID = 100
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = num_part._element

    # Remove any prior definitions with our IDs to keep idempotent
    for old in root.findall(f"{{{W}}}abstractNum[@{{{W}}}abstractNumId='{ABSTRACT_ID}']"):
        root.remove(old)
    for old in root.findall(f"{{{W}}}num[@{{{W}}}numId='{NUM_ID}']"):
        root.remove(old)

    # Build abstractNum
    abstract_xml = f"""<w:abstractNum xmlns:w="{W}" w:abstractNumId="{ABSTRACT_ID}">
  <w:multiLevelType w:val="singleLevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="bullet"/>
    <w:lvlText w:val="•"/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="216" w:hanging="216"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
    </w:rPr>
  </w:lvl>
</w:abstractNum>"""

    num_xml = f"""<w:num xmlns:w="{W}" w:numId="{NUM_ID}">
  <w:abstractNumId w:val="{ABSTRACT_ID}"/>
</w:num>"""

    abstract_elem = etree.fromstring(abstract_xml)
    num_elem = etree.fromstring(num_xml)

    # Insert before any existing <w:num> elements (abstractNum must come first)
    first_num = root.find(f"{{{W}}}num")
    if first_num is not None:
        root.insert(list(root).index(first_num), abstract_elem)
    else:
        root.append(abstract_elem)
    root.append(num_elem)

    # Attach numPr to the style's pPr (use style.element.get_or_add_pPr()
    # because paragraph_format._element returns the CT_Style root for styles)
    pPr = style.element.get_or_add_pPr()
    # Remove existing numPr if any
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)

    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), str(NUM_ID))
    numPr.append(ilvl)
    numPr.append(numId_elem)
    pPr.append(numPr)


def style_skills_line(doc):
    """
    Skills Line — Pipe-separated skills row.
    11pt Calibri, black, 0pt before, 8pt after.
    """
    s = _get_or_add_style(doc, "Skills Line", base_style_name="Normal")
    _set_font(s.font, "Calibri", 11, color=BLACK)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(s.paragraph_format, before_pt=0, after_pt=8)
    return s


def style_accomplishment(doc):
    """
    Accomplishment — Impact bullet identical to List Bullet in formatting.
    11pt Calibri, black, bullet glyph •, 0.15in indent, 2pt before, 2pt after.
    Kept as a separate style so the renderer can distinguish accomplishment
    bullets from plain list bullets semantically.
    """
    s = _get_or_add_style(doc, "Accomplishment", base_style_name="List Bullet")
    # Inherit List Bullet; just confirm font in case base changes
    _set_font(s.font, "Calibri", 11, color=BLACK)
    _set_para_spacing(s.paragraph_format, before_pt=2, after_pt=2)
    return s


def style_body_text_summary(doc):
    """
    Body Text (Summary) — Summary / narrative paragraph.
    11pt Calibri, black, justified, 0pt before, 8pt after.

    Display name: "Body Text (Summary)" — parentheses are accepted by
    python-docx 1.2.0. Word normalises the style ID to "BodyTextSummary"
    internally but the display name is preserved.
    """
    s = _get_or_add_style(doc, "Body Text (Summary)", base_style_name="Normal")
    _set_font(s.font, "Calibri", 11, color=BLACK)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_para_spacing(s.paragraph_format, before_pt=0, after_pt=8)
    return s


# ── Main ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page setup
    configure_page(doc)

    # Apply styles (order matters for built-ins — override after Document() init)
    style_heading1(doc)
    style_heading2(doc)
    style_heading3(doc)
    style_tagline(doc)
    style_contact(doc)
    style_role_meta(doc)
    style_list_bullet(doc)
    style_skills_line(doc)
    style_accomplishment(doc)
    style_body_text_summary(doc)

    # Remove the default empty paragraph that Document() always adds.
    # The template body must be empty (no sample content).
    # We cannot fully remove the last paragraph in docx (Word requires >=1),
    # so we leave the single empty paragraph with no style assignment.
    # It will not appear in the rendered output as visible content.

    doc.save(OUTPUT_PATH)
    print(f"Template written to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
